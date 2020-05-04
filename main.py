# William Escobar
# September 2019
# All Rights Reserved
# Version 1.0

import docx 
import os, sys, subprocess
from docx.shared import Pt
doc = docx.Document('template.docx')
doc2 = docx.Document('template_2.docx')


def function_create_label_1():
    # This creates and saves the INTERIOR shipping label.
    # This parses the input to the shipping document.

    doc.add_paragraph('Project Number: ' + project_num)
    doc.add_paragraph('EB Part Number: ' + eb_pn)
    doc.add_paragraph('TR Part Number: ' + tr_pn)
    doc.add_paragraph('Description: VAL, TYPE, MATERIAL, DIM')
    doc.add_paragraph('Quantity: ' + qty_string())
    doc.add_paragraph('Purchase Order: ' + po_num)
    doc.add_paragraph('Serial Number: ' + concat_serial())

    doc.save('1.docx')

    print('Shipping labels have been created!')

def function_create_label_2():
    # This creates and saves the EXTERIOR shipping label.
    # This parses the input to the shipping document.

    section = doc2.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "TR SHIPPER: " + project_num + "-" + po_num + "\tRef Number: " + ref_num

    doc2.add_paragraph('\t\tPurchase Order: ' + po_num)
    doc2.add_paragraph('\t\tProject Number: ' + project_num)
    doc2.add_paragraph('\t\tV-PAC Packing List Number: ' + vpac)

    doc2.save('2.docx')

    print('Shipping labels have been created!')

def display_input_menu():
    print('You entered:\n')
    print('Ref Number: ' + ref_num)
    print('Project Number: ' + project_num)
    print('Purchase Order: ' + po_num)
    print('TR Part Number: ' + tr_pn)
    print('EB Part Number: ' + eb_pn)
    print('PO Line Item: ' + poli)
    print('\nSelect an option from the menu below:\n\n'
          '1. Create Labels\n'
          '2. Fix EB Reference Number\n'
          '3. Fix EB Project Number\n'
          '4. Fix EB Purchase Order\n'
          '5. Fix TR Part Number\n'
          '6. Fix EB Part Number\n'
          '7. Fix PO Line Item\n'
          '8. Exit\n')

def options_menu():
    print('1. Create Labels\n'
          '2. Fix EB Reference Number\n'
          '3. Fix EB Project Number\n'
          '4. Fix EB Purchase Order\n'
          '5. Fix TR Part Number\n'
          '6. Fix EB Part Number\n'
          '7. Fix Serial Number\n'
          '8. Exit\n')

def update_ref_num():
    global ref_num
    ref_num = input('Update EB Reference Number: ')
    print('\nUpdated EB Reference Number!\n')
def update_project_num():
    global project_num
    project_num = input('Update EB Project Number: ')
    print('\nUpdated EB Project Number!\n')
def update_po_num():
    global po_num
    po_num = input('Update EB Purchase Order: ')
    print('\nUpdated EB Purchase Order!\n')
def update_tr_pn():
    global tr_pn
    tr_pn = input('Update TR Part Number: ')
    print('\nUpdated TR Part Number!\n')
def update_eb_pn():
    global eb_pn
    eb_pn = input('Update EB Part Number: ')
    print('\nUpdated EB Part Number!\n')
def update_poli():
    global poli
    poli = input('Update Serial Number: ')
    print('\nUpdated Serial Number!\n')
def update_qty():
    global qty
    qty = int(input('Update POLI Quantity: '))
    print('\nUpdated POLI Quantity!\n')
def concat_serial():
    global serialnumber
    serialnumber = (po_num + '-' + str(poli) + '-' + str(qty))
    return serialnumber
def qty_string ():
    global qty
    qty = str(qty)
    return qty

# Main Menu
print('Enter the information as prompted:\n')
ref_num = input('EB Reference Number: ')
project_num = input('EB Project Number: ')
po_num = input('EB Purchase Order: ')
tr_pn = input('TR Part Number: ')
eb_pn = input('EB Part Number: ')
poli = input('PO Line Item: ')
qty = int(input('Quantity: '))
vpac = input('VPAC Number: ')
display_input_menu()


while True:
    try:
        menu_selection = int(input('Enter option: '))

        if menu_selection == 1:
            function_create_label_1()
            function_create_label_2()

            if sys.platform == "win32":
                os.startfile('1.docx')
                os.startfile('2.docx')
            else:
                opener = "open" if sys.platform == "darwin" else "xdg-open"
                subprocess.call([opener, '1.docx'])
                subprocess.call([opener, '2.docx'])
            break

        elif menu_selection == 2:
            update_ref_num()
            display_input_menu()

        elif menu_selection == 3:
            update_project_num()
            display_input_menu()

        elif menu_selection == 4:
            update_po_num()
            display_input_menu()

        elif menu_selection == 5:
            update_tr_pn()
            display_input_menu()

        elif menu_selection == 6:
            update_eb_pn()
            display_input_menu()

        elif menu_selection == 7:
            update_poli()
            display_input_menu()

        elif menu_selection == 8:
            update_qty()
            display_input_menu()

        elif menu_selection == 9:
            break

        else:
            print('\nInvalid choice. Enter one of the options below.\n')
            options_menu()
    except ValueError:
        print('\nInvalid choice. Enter one of the options below.\n')
        options_menu()
exit